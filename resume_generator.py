from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def set_font(run, size=12, bold=False):
    run.font.name = 'Garamond'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Garamond')
    run.font.size = Pt(size)
    run.bold = bold

def generate_resume_docx(user_data, photo_path=None):
    doc = Document()

    # 1x2 jadval: chap (matn), o‘ng (rasm)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(4.5)
    table.columns[1].width = Inches(2.0)

    # Chap qism (asosiy matnlar)
    cell_left = table.cell(0, 0)

    def add_paragraph_to_cell(cell, text, size=12, bold=False):
        para = cell.add_paragraph()
        run = para.add_run(text)
        set_font(run, size=size, bold=bold)

    # Ism
    add_paragraph_to_cell(cell_left, user_data['name'], size=22, bold=True)

    # Aloqa
    add_paragraph_to_cell(cell_left, f"📧 {user_data['email']} | 📞 {user_data['phone']}", size=10)

    # 🎯 Maqsad
    add_paragraph_to_cell(cell_left, "\n🎯 Maqsad", size=14, bold=True)
    add_paragraph_to_cell(cell_left, user_data['objective'])

    # 🎓 Ta’lim
    add_paragraph_to_cell(cell_left, "\n🎓 Ta’lim", size=14, bold=True)
    for line in user_data['education'].split('\n'):
        add_paragraph_to_cell(cell_left, line)

    # 💼 Ish tajribasi
    add_paragraph_to_cell(cell_left, "\n💼 Ish tajribasi", size=14, bold=True)
    for line in user_data['experience'].split('\n'):
        parts = [p.strip() for p in line.split(',')]
        if len(parts) == 3:
            company, years, role = parts
            add_paragraph_to_cell(cell_left, f"🏢 {company} ({years})", bold=True)
            add_paragraph_to_cell(cell_left, role)

    # 🛠 Ko‘nikmalar
    add_paragraph_to_cell(cell_left, "\n🛠 Ko‘nikmalar", size=14, bold=True)
    for skill in user_data['skills'].split(','):
        add_paragraph_to_cell(cell_left, f"• {skill.strip()}")

    # 🌐 Tillar
    add_paragraph_to_cell(cell_left, "\n🌐 Tillar", size=14, bold=True)
    for lang in user_data['languages'].split('\n'):
        add_paragraph_to_cell(cell_left, f"• {lang.strip()}")

    # 📊 Office dasturlari
    add_paragraph_to_cell(cell_left, "\n📊 Office dasturlari", size=14, bold=True)
    for off in user_data['office'].split(','):
        add_paragraph_to_cell(cell_left, f"• {off.strip()}")

    # Rasmni o‘ng tomonga qo‘shish
    if photo_path and os.path.exists(photo_path):
        cell_right = table.cell(0, 1)
        cell_right.paragraphs[0].add_run().add_picture(photo_path, width=Inches(2))
        cell_right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Faylni saqlash
    filename = f"{user_data['name'].replace(' ', '_')}_resume.docx"
    doc.save(filename)
    return filename
